"""Temp file management for order deliverables.

``DeliverableManager`` organises generated files into per-order
directories under a configurable base path and provides convenience
methods for the most common deliverable formats (plain text, Word, Excel,
Python source).

Usage::

    from src.utils.file_handler import DeliverableManager

    dm = DeliverableManager()
    path = dm.save_text("order_abc123", "notes.txt", "Hello, world!")
    print(path)  # data/deliverables/order_abc123/notes.txt
"""

from __future__ import annotations

import shutil
from pathlib import Path
from typing import Any

import structlog
from docx import Document as DocxDocument  # python-docx
from openpyxl import Workbook

log = structlog.stdlib.get_logger(__name__)

# Resolve the project root so that the default base_dir is always
# relative to the repository, regardless of the caller's cwd.
_PROJECT_ROOT = Path(__file__).resolve().parents[2]


class DeliverableManager:
    """Create, store, list, and clean up order deliverables."""

    def __init__(self, base_dir: str = "data/deliverables") -> None:
        # If the caller passes a relative path we anchor it to the project
        # root; absolute paths are used as-is.
        path = Path(base_dir)
        if not path.is_absolute():
            path = _PROJECT_ROOT / path
        self._base_dir = path
        self._base_dir.mkdir(parents=True, exist_ok=True)
        log.debug("deliverable_manager.init", base_dir=str(self._base_dir))

    # ------------------------------------------------------------------
    # Path helpers
    # ------------------------------------------------------------------

    def _order_dir(self, order_id: str) -> Path:
        """Return (and create) the directory for *order_id*."""
        d = self._base_dir / order_id
        d.mkdir(parents=True, exist_ok=True)
        return d

    def create_path(self, order_id: str, filename: str) -> Path:
        """Create the order sub-directory and return the full file path.

        The file itself is **not** created -- only the parent directory.
        """
        return self._order_dir(order_id) / filename

    # ------------------------------------------------------------------
    # Writers
    # ------------------------------------------------------------------

    def save_text(self, order_id: str, filename: str, content: str) -> Path:
        """Write *content* to a UTF-8 text file and return its path."""
        path = self.create_path(order_id, filename)
        path.write_text(content, encoding="utf-8")
        log.info(
            "deliverable.saved",
            order_id=order_id,
            filename=filename,
            fmt="text",
            size=len(content),
        )
        return path

    def save_docx(
        self,
        order_id: str,
        filename: str,
        content: str,
        title: str = "",
    ) -> Path:
        """Create a Word document with *content* and return its path.

        If *title* is provided it is added as a ``Heading 1`` before the
        body text.  The body is split on newlines so each paragraph in
        the source string becomes a separate Word paragraph.
        """
        if not filename.lower().endswith(".docx"):
            filename = f"{filename}.docx"

        path = self.create_path(order_id, filename)

        doc = DocxDocument()
        if title:
            doc.add_heading(title, level=1)
        for paragraph in content.split("\n"):
            doc.add_paragraph(paragraph)
        doc.save(str(path))

        log.info(
            "deliverable.saved",
            order_id=order_id,
            filename=filename,
            fmt="docx",
            size=path.stat().st_size,
        )
        return path

    def save_xlsx(
        self,
        order_id: str,
        filename: str,
        data: list[list[Any]],
        headers: list[str] | None = None,
    ) -> Path:
        """Create an Excel workbook and return its path.

        Parameters
        ----------
        data:
            A list of rows, where each row is a list of cell values.
        headers:
            Optional column headers written as the first row.
        """
        if not filename.lower().endswith(".xlsx"):
            filename = f"{filename}.xlsx"

        path = self.create_path(order_id, filename)

        wb = Workbook()
        ws = wb.active
        if ws is None:  # Defensive; should never happen with a new Workbook.
            ws = wb.create_sheet()

        if headers:
            ws.append(headers)
        for row in data:
            ws.append(row)
        wb.save(str(path))

        log.info(
            "deliverable.saved",
            order_id=order_id,
            filename=filename,
            fmt="xlsx",
            rows=len(data),
            size=path.stat().st_size,
        )
        return path

    def save_python(self, order_id: str, filename: str, code: str) -> Path:
        """Save a Python source file after validating syntax.

        The code is compiled with :func:`compile` to catch syntax errors
        *before* persisting, so that callers receive an immediate
        ``SyntaxError`` rather than a silently broken file.
        """
        if not filename.lower().endswith(".py"):
            filename = f"{filename}.py"

        # Validate syntax.  compile() raises SyntaxError on invalid code.
        compile(code, filename, "exec")

        path = self.create_path(order_id, filename)
        path.write_text(code, encoding="utf-8")

        log.info(
            "deliverable.saved",
            order_id=order_id,
            filename=filename,
            fmt="python",
            size=len(code),
        )
        return path

    # ------------------------------------------------------------------
    # Query / cleanup
    # ------------------------------------------------------------------

    def get_deliverables(self, order_id: str) -> list[Path]:
        """Return a sorted list of all files belonging to *order_id*.

        Returns an empty list if the order directory does not exist.
        """
        order_dir = self._base_dir / order_id
        if not order_dir.is_dir():
            return []
        return sorted(p for p in order_dir.iterdir() if p.is_file())

    def cleanup(self, order_id: str) -> None:
        """Recursively remove the order directory and all its contents."""
        order_dir = self._base_dir / order_id
        if order_dir.is_dir():
            shutil.rmtree(order_dir)
            log.info("deliverable.cleanup", order_id=order_id)
        else:
            log.debug(
                "deliverable.cleanup.noop",
                order_id=order_id,
                reason="directory does not exist",
            )
