"""Tests for the deliverable file manager (``src.utils.file_handler``)."""

from __future__ import annotations

from pathlib import Path

import pytest

from src.utils.file_handler import DeliverableManager


# =========================================================================
# Fixtures
# =========================================================================

ORDER_ID = "order-file-test"


@pytest.fixture
def dm(tmp_path: Path) -> DeliverableManager:
    """Return a DeliverableManager rooted in a pytest temporary directory."""
    return DeliverableManager(base_dir=str(tmp_path / "deliverables"))


# =========================================================================
# save_text
# =========================================================================


class TestSaveText:
    """Verify save_text writes correct content to disk."""

    def test_creates_file(self, dm: DeliverableManager) -> None:
        path = dm.save_text(ORDER_ID, "notes.txt", "Hello, world!")
        assert path.exists()
        assert path.name == "notes.txt"

    def test_content_matches(self, dm: DeliverableManager) -> None:
        content = "Line 1\nLine 2\nLine 3"
        path = dm.save_text(ORDER_ID, "content.txt", content)
        assert path.read_text(encoding="utf-8") == content

    def test_unicode_content(self, dm: DeliverableManager) -> None:
        content = "Cafe\u0301 au lait -- \u2603 snowman"
        path = dm.save_text(ORDER_ID, "unicode.txt", content)
        assert path.read_text(encoding="utf-8") == content

    def test_empty_content(self, dm: DeliverableManager) -> None:
        path = dm.save_text(ORDER_ID, "empty.txt", "")
        assert path.exists()
        assert path.read_text(encoding="utf-8") == ""

    def test_order_directory_created(self, dm: DeliverableManager) -> None:
        path = dm.save_text(ORDER_ID, "test.txt", "data")
        assert path.parent.name == ORDER_ID
        assert path.parent.is_dir()


# =========================================================================
# save_docx
# =========================================================================


class TestSaveDocx:
    """Verify save_docx creates a valid Word document."""

    def test_creates_file(self, dm: DeliverableManager) -> None:
        path = dm.save_docx(ORDER_ID, "article.docx", "Some content")
        assert path.exists()
        assert path.suffix == ".docx"

    def test_appends_extension(self, dm: DeliverableManager) -> None:
        """If the filename lacks .docx the method should append it."""
        path = dm.save_docx(ORDER_ID, "article", "Content")
        assert path.name == "article.docx"

    def test_content_readable(self, dm: DeliverableManager) -> None:
        """Open the file with python-docx and verify body text is present."""
        from docx import Document as DocxDocument

        text = "Paragraph one.\nParagraph two."
        path = dm.save_docx(ORDER_ID, "doc.docx", text)
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        assert "Paragraph one." in paragraphs
        assert "Paragraph two." in paragraphs

    def test_title_heading(self, dm: DeliverableManager) -> None:
        """When a title is provided it should appear as a heading."""
        from docx import Document as DocxDocument

        path = dm.save_docx(ORDER_ID, "titled.docx", "Body text", title="My Title")
        doc = DocxDocument(str(path))
        # The first paragraph with style containing "Heading" holds the title
        first_para = doc.paragraphs[0]
        assert first_para.text == "My Title"

    def test_file_is_valid_zip(self, dm: DeliverableManager) -> None:
        """A .docx is technically a ZIP archive; verify it is valid."""
        import zipfile

        path = dm.save_docx(ORDER_ID, "zip_check.docx", "Content here")
        assert zipfile.is_zipfile(str(path))


# =========================================================================
# save_python
# =========================================================================


class TestSavePython:
    """Verify save_python creates a syntactically valid .py file."""

    def test_creates_file(self, dm: DeliverableManager) -> None:
        code = "def greet():\n    return 'hello'\n"
        path = dm.save_python(ORDER_ID, "script.py", code)
        assert path.exists()
        assert path.suffix == ".py"

    def test_appends_extension(self, dm: DeliverableManager) -> None:
        path = dm.save_python(ORDER_ID, "utils", "x = 1\n")
        assert path.name == "utils.py"

    def test_content_matches(self, dm: DeliverableManager) -> None:
        code = "print('hello')\n"
        path = dm.save_python(ORDER_ID, "hello.py", code)
        assert path.read_text(encoding="utf-8") == code

    def test_syntax_validation_passes(self, dm: DeliverableManager) -> None:
        code = (
            "import os\n"
            "\n"
            "def main() -> None:\n"
            "    for i in range(10):\n"
            "        print(os.getpid(), i)\n"
        )
        path = dm.save_python(ORDER_ID, "valid.py", code)
        assert path.exists()

    def test_invalid_syntax_raises(self, dm: DeliverableManager) -> None:
        bad_code = "def broken(:\n    pass\n"
        with pytest.raises(SyntaxError):
            dm.save_python(ORDER_ID, "bad.py", bad_code)

    def test_invalid_syntax_does_not_create_file(
        self, dm: DeliverableManager
    ) -> None:
        bad_code = "class Foo(\n"
        with pytest.raises(SyntaxError):
            dm.save_python(ORDER_ID, "no_file.py", bad_code)
        # The file should not have been written
        deliverables = dm.get_deliverables(ORDER_ID)
        file_names = [p.name for p in deliverables]
        assert "no_file.py" not in file_names


# =========================================================================
# save_xlsx
# =========================================================================


class TestSaveXlsx:
    """Verify save_xlsx creates a valid Excel workbook."""

    def test_creates_file(self, dm: DeliverableManager) -> None:
        data = [["Alice", 30], ["Bob", 25]]
        path = dm.save_xlsx(ORDER_ID, "people.xlsx", data)
        assert path.exists()
        assert path.suffix == ".xlsx"

    def test_appends_extension(self, dm: DeliverableManager) -> None:
        path = dm.save_xlsx(ORDER_ID, "data", [["row"]])
        assert path.name == "data.xlsx"

    def test_content_readable(self, dm: DeliverableManager) -> None:
        from openpyxl import load_workbook

        data = [["Alice", 30], ["Bob", 25]]
        headers = ["Name", "Age"]
        path = dm.save_xlsx(ORDER_ID, "table.xlsx", data, headers=headers)

        wb = load_workbook(str(path))
        ws = wb.active
        assert ws is not None

        # Row 1: headers
        assert ws.cell(1, 1).value == "Name"
        assert ws.cell(1, 2).value == "Age"
        # Row 2: first data row
        assert ws.cell(2, 1).value == "Alice"
        assert ws.cell(2, 2).value == 30
        # Row 3: second data row
        assert ws.cell(3, 1).value == "Bob"
        assert ws.cell(3, 2).value == 25

    def test_without_headers(self, dm: DeliverableManager) -> None:
        from openpyxl import load_workbook

        data = [["X", 1], ["Y", 2]]
        path = dm.save_xlsx(ORDER_ID, "no_headers.xlsx", data)

        wb = load_workbook(str(path))
        ws = wb.active
        assert ws is not None
        # First row should be data, not headers
        assert ws.cell(1, 1).value == "X"

    def test_file_is_valid_zip(self, dm: DeliverableManager) -> None:
        import zipfile

        path = dm.save_xlsx(ORDER_ID, "zip_check.xlsx", [["val"]])
        assert zipfile.is_zipfile(str(path))


# =========================================================================
# get_deliverables
# =========================================================================


class TestGetDeliverables:
    """Verify get_deliverables lists files correctly."""

    def test_lists_created_files(self, dm: DeliverableManager) -> None:
        dm.save_text(ORDER_ID, "a.txt", "aaa")
        dm.save_text(ORDER_ID, "b.txt", "bbb")
        files = dm.get_deliverables(ORDER_ID)
        names = [p.name for p in files]
        assert "a.txt" in names
        assert "b.txt" in names
        assert len(files) == 2

    def test_returns_sorted(self, dm: DeliverableManager) -> None:
        dm.save_text(ORDER_ID, "z.txt", "z")
        dm.save_text(ORDER_ID, "a.txt", "a")
        dm.save_text(ORDER_ID, "m.txt", "m")
        files = dm.get_deliverables(ORDER_ID)
        names = [p.name for p in files]
        assert names == sorted(names)

    def test_empty_for_nonexistent_order(self, dm: DeliverableManager) -> None:
        files = dm.get_deliverables("no-such-order")
        assert files == []

    def test_returns_paths(self, dm: DeliverableManager) -> None:
        dm.save_text(ORDER_ID, "file.txt", "data")
        files = dm.get_deliverables(ORDER_ID)
        assert all(isinstance(p, Path) for p in files)

    def test_mixed_file_types(self, dm: DeliverableManager) -> None:
        dm.save_text(ORDER_ID, "readme.txt", "text")
        dm.save_docx(ORDER_ID, "doc.docx", "content")
        dm.save_python(ORDER_ID, "script.py", "x = 1\n")
        dm.save_xlsx(ORDER_ID, "data.xlsx", [["a"]])
        files = dm.get_deliverables(ORDER_ID)
        assert len(files) == 4


# =========================================================================
# cleanup
# =========================================================================


class TestCleanup:
    """Verify cleanup removes the order directory and all files."""

    def test_removes_directory(self, dm: DeliverableManager) -> None:
        dm.save_text(ORDER_ID, "file.txt", "data")
        assert len(dm.get_deliverables(ORDER_ID)) == 1

        dm.cleanup(ORDER_ID)
        assert dm.get_deliverables(ORDER_ID) == []

    def test_cleanup_nonexistent_is_noop(self, dm: DeliverableManager) -> None:
        """Cleaning up a non-existent order should not raise."""
        dm.cleanup("does-not-exist")  # no exception

    def test_cleanup_removes_all_files(self, dm: DeliverableManager) -> None:
        dm.save_text(ORDER_ID, "a.txt", "a")
        dm.save_text(ORDER_ID, "b.txt", "b")
        dm.save_docx(ORDER_ID, "c.docx", "c")
        assert len(dm.get_deliverables(ORDER_ID)) == 3

        dm.cleanup(ORDER_ID)
        assert dm.get_deliverables(ORDER_ID) == []
